from transformers import AlbertTokenizer, AlbertTokenizerFast
from transcription import Transcription, Transcription_manager
import os
from docx import Document
import mammoth
import contractions
import re


doc_file_path =  None #r"C:\Users\12058\Documents\AR_QP0I--R_Qp0idSrCxm7OCat_Mirielle.docx"

ghost_transcript = ""
clown_transcript = ""

ghost_cotractionless = ""
clown_contractionless = ""

tokens_ghost = None
tokens_ghost_count = 0

tokens_clown = None
tokens_clown_count = 0

def is_bold(run):
    return run.bold

def process(doc_file_path, transcription, collection):

    if not os.path.exists(doc_file_path):
        print(f"Error: File '{doc_file_path}' does not exist.")
    elif os.path.getsize(doc_file_path) == 0:
        print(f"Error: File '{doc_file_path}' is empty.")
    else:
        doc = Document(doc_file_path)

        full_text = "\n".join([p.text for p in doc.paragraphs])
        transcription.set_transcript_ID(full_text)

        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if(is_bold(run)):
                    run.text = ""
        doc.save(doc_file_path)

        with open(doc_file_path, "rb") as docx_file:
            result = mammoth.extract_raw_text(docx_file)
            text = result.value
            #print(text)

        paragraphs = [segment.strip() for segment in re.split(r'\n{2,}', text) if segment.strip()]
        ghost_transcript = paragraphs[0]
        clown_transcript = paragraphs[1]
        # print("ghost transcript: ", ghost_transcript, "\n")
        # print("clown _transcript: ", clown_transcript, "\n")


        ghost_contractionless = re.sub(r'\[incoherent\]|\[inaudible\]', '', contractions.fix(ghost_transcript))
        clown_contractionless = re.sub(r'\[incoherent\]|\[inaudible\]', '', contractions.fix(clown_transcript))
        transcription.set_cleaned_text(clown_contractionless, ghost_contractionless) #TEST THIS
        collection.add_transcription(transcription)

        #print("ghost transcript no contractions: ", ghost_contractionless, "\n")
        #print("clown transcript no contractions: ", clown_contractionless, "\n")




        # tokenizer = AlbertTokenizer.from_pretrained('albert-base-v2')
        # inputs = tokenizer(ghost_contractionless, return_tensors='pt')
        # input_ids = inputs['input_ids']
        # tokens_ghost_count = len(inputs['input_ids'])
        # tokens_ghost = tokenizer.convert_ids_to_tokens(input_ids[0])
        # print(tokens_ghost, "\n")
        # print(tokens_ghost_count)

        # inputs = tokenizer(clown_contractionless, return_tensors='pt')
        # input_ids = inputs['input_ids']
        # tokens_clown_count = len(inputs['input_ids'])
        # tokens_clown = tokenizer.convert_ids_to_tokens(input_ids[0])
        # print(tokens_clown, "\n")
        # print(tokens_clown_count)